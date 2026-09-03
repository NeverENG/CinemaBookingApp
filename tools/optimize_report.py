# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = '/Users/4ge0/Downloads/9110125188-吕宇轩-电影院订票系统-原格式补充版_副本2.docx'
OUT = '/Users/4ge0/Desktop/code/HomeWork/LTerm/9110125188-吕宇轩-电影院订票系统-优化版.docx'
ASSET_DIR = '/Users/4ge0/Desktop/code/HomeWork/LTerm/report-assets'

doc = Document(SRC)

# Use a Chinese font available both on this Mac and in common WPS installs;
# the original template names Songti/SimSun, which is absent in this runtime.
for style_name in ['Normal', 'List Paragraph']:
    try:
        st = doc.styles[style_name]
        st.font.name = 'STSong'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'STSong')
    except KeyError:
        pass
for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = 'STSong'
        if r._element.rPr is not None:
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'STSong')

def clear_para(p):
    for child in list(p._p):
        if child.tag != qn('w:pPr'):
            p._p.remove(child)

def set_text(p, text, bold=False):
    clear_para(p)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'STSong'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'STSong')
    r.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def set_code(p, code):
    clear_para(p)
    p.style = doc.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    pf.left_indent = Inches(0.12)
    pf.right_indent = Inches(0.12)
    ppr = p._p.get_or_add_pPr()
    for tag in ('pStyle', 'numPr', 'keepLines', 'keepNext', 'widowControl'):
        node = ppr.find(qn('w:' + tag))
        if node is not None:
            ppr.remove(node)
    # WPS renders this paragraph as an editable code block: light gray fill
    # and a thin neutral border, while preserving copyable line breaks.
    ppr = p._p.get_or_add_pPr()
    shd = ppr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); ppr.append(shd)
    shd.set(qn('w:fill'), 'F2F2F2')
    borders = ppr.find(qn('w:pBdr'))
    if borders is None:
        borders = OxmlElement('w:pBdr'); ppr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = qn('w:' + edge)
        el = borders.find(tag)
        if el is None:
            el = OxmlElement('w:' + edge); borders.append(el)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '3'); el.set(qn('w:color'), 'BFBFBF')
    for n, line in enumerate(code.splitlines()):
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        r.font.size = Pt(9)
        if n < len(code.splitlines()) - 1:
            r.add_break()

def insert_after(p, text, style=None):
    new_p = OxmlElement('w:p')
    p._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    np = Paragraph(new_p, p._parent)
    if style:
        np.style = style
    set_text(np, text)
    return np

def insert_figure_after(anchor, image_path, caption, lead):
    lead_p = insert_after(anchor, lead)
    lead_p.paragraph_format.keep_with_next = True
    image_p = insert_after(lead_p, '')
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.keep_with_next = False
    image_p.paragraph_format.keep_together = False
    run = image_p.add_run()
    run.add_picture(image_path, width=Inches(5.35))
    caption_p = insert_after(image_p, caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in caption_p.runs:
        r.font.name = 'STSong'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'STSong')
        r.font.size = Pt(10.5)
    return caption_p

def find_paragraph(text):
    for p in doc.paragraphs:
        if text in p.text:
            return p
    raise ValueError(f'paragraph not found: {text}')

def remove_image(p):
    for node in list(p._p.xpath('.//w:drawing')):
        node.getparent().remove(node)

# Remove report screenshots/diagram screenshots while retaining cover artwork.
for idx in [82, 84, 97, 108, 110, 112, 118]:
    remove_image(doc.paragraphs[idx])

# Roles: use the existing blank paragraphs in the 2.3 section.
roles = [
    '角色职责与权限边界：平台管理员面向平台级资源与权限治理，负责接收并维护片方提供的影片资料，维护影院基础信息，审核影片并控制上下架，配置账号、角色和权限；当影院停业、影片下线或账号异常时，可执行禁用、恢复及审计追踪。平台管理员不直接修改已支付订单和支付流水。',
    '影院管理员只管理被授权影院，负责影厅与座位维护、排片创建与调整、票价配置、场次关闭以及销售情况查看。取消场次时，系统必须释放相关锁座并通知用户；影院管理员不得修改平台角色权限和财务流水。',
    '财务人员围绕资金与经营数据工作，负责支付流水核对、退款审核、票房统计和日报对账，不直接修改订单业务状态。普通用户完成注册登录、影片/影院/场次浏览、实时选座、锁座购票、订单查询、退票及积分优惠券使用。典型链路为“管理员录入影片→审核上架→关联影院排片；用户浏览影片→选择影院和场次→锁座→创建订单→支付回调出票；财务按支付、退款流水与日票房数据对账”。'
]
for idx, txt in zip([68, 69, 70], roles):
    set_text(doc.paragraphs[idx], txt)
# Move role details to the beginning of section 2.3, before numbered modules.
role_paras = [doc.paragraphs[i] for i in [68, 69, 70]]
role_anchor = doc.paragraphs[62]
for p in role_paras:
    role_anchor._p.addnext(p._p)
    role_anchor = p

# The original architecture screenshot location was after all feature details.
# Clear that obsolete placeholder; the redrawn overview is placed at the start
# of section 2.2 so readers see the whole system before module details.
set_text(doc.paragraphs[81], '')
set_text(doc.paragraphs[82], '')

# Database section: dedicated classification and flow prose.
set_text(doc.paragraphs[84], '数据库表归类与数据流转关系')
db_flow = ('本系统按照业务边界将数据表划分为六类：账户与会员域（users、membership_levels、membership_level_logs、points_ledger）记录身份、等级和积分流水；影片与影院资源域（movies、cinemas、halls、seats）描述影片、影院、影厅和物理座位；排片域（show_sessions）描述影片在哪家影院、哪个影厅及何时上映；交易订单域（seat_locks、orders、order_items）承载选座、锁座、订单和电子票信息；支付售后域（payment_transactions、payment_callbacks、refunds）记录支付回调、退款和资金状态；营销、统计与审计域（coupon_templates、user_coupons、box_office_ledger、daily_box_office、operation_logs）支撑优惠、票房汇总和操作追踪。数据从主数据开始流动：管理员维护 movies、cinemas、halls 和 seats 后创建 show_sessions；用户读取影片、影院、场次和座位信息，并在 seat_locks 中以 (session_id, seat_id) 部分唯一索引抢占座位；锁座成功后，系统在同一事务内创建 orders 和 order_items，订单创建时生成 ticket_no，支付成功后才允许使用和核销，随后创建 payment_transactions。支付回调先以 event_id 幂等写入 payment_callbacks，再将订单由 PENDING_PAYMENT 更新为 PAID、锁座由 LOCKED 更新为 BOOKED，使订单明细中的电子票进入可使用状态，同时记录积分流水和票房事件，并更新 daily_box_office 聚合数据。取消或超时使订单进入 CANCELED 或 EXPIRED 状态并释放锁座；退款则联动 orders、payment_transactions、seat_locks、points_ledger、box_office_ledger 和 daily_box_office 完成状态回写。各表通过主键、外键和业务单号保持引用完整；支付与退款状态表允许按状态机更新，payment_callbacks、points_ledger、box_office_ledger 和 operation_logs 等事件流水以追加方式保留审计记录。')
set_text(doc.paragraphs[85], db_flow)
insert_after(find_paragraph('2.1 总体设计'), '系统设计以“资源维护—场次售卖—订单履约—售后对账”为主线。影片、影院和座位属于相对稳定的主数据，场次是可售资源的组合，订单和支付记录则保存一次购票行为的完整生命周期。这样划分既便于管理员按职责维护数据，也便于在退款、对账和审计时追溯原始业务事件。')

db_table = doc.add_table(rows=1, cols=4)
db_table.style = 'Table Grid'
hdr = db_table.rows[0].cells
for c, t in zip(hdr, ['数据归类', '核心表', '主要职责', '关键关联']):
    c.text = t
rows = [
    ('账户与会员', 'users、membership_levels、points_ledger', '身份、权限、等级、积分流水', 'user_id；(biz_type,biz_no)唯一'),
    ('影片与影院资源', 'movies、cinemas、halls、seats', '影片、影院、影厅、物理座位', 'cinema_id、hall_id、movie_id'),
    ('排片', 'show_sessions', '影片、影院、影厅、时间与票价', 'cinema_id、hall_id、movie_id'),
    ('交易订单', 'seat_locks、orders、order_items', '锁座、订单、电子票明细', 'session_id、seat_id、order_no'),
    ('支付与售后', 'payment_transactions、payment_callbacks、refunds', '支付回调、退款、资金状态', 'biz_no、event_id、order_id、refund_no'),
    ('营销统计审计', 'user_coupons、box_office_ledger、daily_box_office、operation_logs', '券核销、票房流水、看板与操作留痕', 'order_no、cinema_id、movie_id'),
]
for row in rows:
    cells = db_table.add_row().cells
    for c, t in zip(cells, row):
        c.text = t
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = 'STSong'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'STSong'); r.font.size = Pt(10)
# python-docx appends new tables to the document body; relocate this table to
# the database section so it does not create a standalone trailing page.
tbl = db_table._element
doc.paragraphs[85]._p.addnext(tbl)
insert_after(doc.paragraphs[85], '表2-1 核心数据库表归类、职责与关联')
insert_after(doc.paragraphs[85], '主要幂等项与并发控制：order_no、transaction_no、refund_no 和 payment_callbacks.event_id 均设置唯一约束；points_ledger 使用 UNIQUE(biz_type, biz_no) 防止重复赠送或扣回；refunds.order_id UNIQUE 保证一单一退；seat_locks 建立 WHERE status IN (\'LOCKED\', \'BOOKED\') 的部分唯一索引防止超卖；状态迁移使用“WHERE status = 旧状态 AND version = 旧版本”的条件 UPDATE，影响行数为 0 时判定为重复请求或并发失败。')
insert_after(find_paragraph('主要幂等项与并发控制：'), '在异常场景下，系统将重复提交、回调重试、任务重复执行和并发抢座视为正常输入。每个入口通过业务单号、事件号或版本号识别请求，成功结果可以重复读取，失败结果不会留下半成品数据。')

# Replace state-machine screenshots with copyable code and explanations.
set_text(doc.paragraphs[107], '订单与支付状态机：')
set_text(doc.paragraphs[108], '代码5-1 支付回调事务中的核心状态迁移（源码节选）')
set_text(doc.paragraphs[109], '')
set_code(doc.paragraphs[110], '''if order.PaidCents != in.AmountCents {
    return domain.ErrPaymentAmountMismatch
}
if err := payment.Transition(domain.PaymentEventSuccess); err != nil {
    return err
}
if err := s.payments.Transition(txCtx, payment.TransactionNo,
    domain.PaymentPending, domain.PaymentSuccess, payment.Version); err != nil {
    return err
}
if err := order.Transition(domain.OrderEventPaySuccess); err != nil {
    return err
}
if err := s.orders.Transition(txCtx, order.OrderNo,
    domain.OrderPendingPayment, domain.OrderPaid, order.Version); err != nil {
    return err
}
if err := s.locks.MarkBookedByOrderNo(txCtx, order.OrderNo); err != nil {
    return err
}
if err := s.points.GrantOnPaid(txCtx, order.UserID,
    order.PaidCents, order.OrderNo); err != nil {
    return err
}
// 以下省略优惠券核销、会员升级和票房事件写入。
return s.callbacks.MarkProcessed(txCtx, cb.EventID)''')
set_text(doc.paragraphs[112], '代码说明：输入为支付事件 event_id、订单号和金额。系统先检查回调事件是否已经处理，再校验订单状态和支付金额；Repository 在执行状态迁移时校验旧状态与版本。订单、座位、优惠券、积分和票房写入处于同一事务，任一步骤失败都会整体回滚，因此重复回调不会重复出票、记账或增加积分。')
doc.paragraphs[110]._p.addnext(doc.paragraphs[112]._p)

# Add concise, copyable excerpts for the other high-risk business functions.
def add_code_excerpt(anchor_text, title, code, explanation):
    anchor = find_paragraph(anchor_text)
    title_p = insert_after(anchor, title)
    title_p.paragraph_format.keep_with_next = True
    code_p = insert_after(title_p, '')
    set_code(code_p, code)
    code_p.paragraph_format.keep_together = True
    note_p = insert_after(code_p, explanation)
    return note_p

add_code_excerpt(
    '5.1，核心请求链路',
    '代码5-2 锁座与创建订单（order_svc.go，源码节选）',
    '''return s.tx.Run(ctx, func(txCtx context.Context) error {
    session, err := s.sessions.GetSessionForUpdate(txCtx, in.SessionID)
    if err != nil || !session.Saleable() { return domain.ErrSessionUnavailable }
    seats, err := s.seats.ListByIDs(txCtx, in.SeatIDs)
    if err != nil { return err }
    if err := s.locks.ReleaseExpiredBySeats(txCtx, in.SessionID, in.SeatIDs); err != nil {
        return err
    }
    total, items, err := s.calculateAmount(session, seats, in.CouponNo)
    if err != nil { return err }
    order = &domain.Order{OrderNo: newOrderNo(), UserID: in.UserID,
        SessionID: in.SessionID, Status: domain.OrderPendingPayment,
        PaidCents: total, ExpireAt: time.Now().Add(orderTTL)}
    order.Items = items
    if err := s.orders.CreateOrder(txCtx, order); err != nil { return err }
    if err := s.locks.CreateLocks(txCtx, buildLocks(order)); err != nil { return err }
    return s.sessions.RecalcStatus(txCtx, order.SessionID)
})''',
    '说明：请求前置条件是场次可售且座位属于该影厅。事务内先锁定场次并清理过期锁，再计算价格、创建待支付订单和 LOCKED 座位锁；任一步失败整体回滚，15 分钟过期时间与活动锁唯一约束共同防止并发超卖。'
)
add_code_excerpt(
    '代码5-2 锁座与创建订单',
    '代码5-3 超时订单释放（order_svc.go，源码节选）',
    '''expired, err := s.orders.ListExpiredPending(ctx, now)
if err != nil { return 0, err }
for _, candidate := range expired {
    err := s.tx.Run(ctx, func(txCtx context.Context) error {
        order, err := s.orders.GetOrderByNo(txCtx, candidate.OrderNo)
        if err != nil || order.Status != domain.OrderPendingPayment { return nil }
        if err := order.Transition(domain.OrderEventTimeout); err != nil { return err }
        if err := s.orders.Transition(txCtx, order.OrderNo,
            domain.OrderPendingPayment, domain.OrderExpired, order.Version); err != nil { return err }
        if err := s.locks.ReleaseByOrderNo(txCtx, order.OrderNo, domain.SeatLockExpired); err != nil { return err }
        if err := s.sessions.RecalcStatus(txCtx, order.SessionID); err != nil { return err }
        return s.coupons.UnlockByOrderNo(txCtx, order.OrderNo)
    })
    if err == nil { count++ }
}
return count, nil''',
    '说明：定时任务先批量扫描过期待支付订单，再在事务内二次读取并检查状态。只有仍为 PENDING_PAYMENT 的订单才通过版本条件更新为 EXPIRED，并同步释放座位、重算场次和解锁优惠券；已支付或已取消订单会被跳过。'
)
add_code_excerpt(
    '代码5-3 超时订单释放',
    '代码5-4 退款回调与资源回收（refund_svc.go，源码节选）',
    '''if refund.Status == domain.RefundSuccess { return nil }
if refund.Status != domain.RefundPending { return domain.ErrRefundState }
order, err := s.orders.GetOrderByNo(txCtx, refund.OrderNo)
if err != nil || order.Status != domain.OrderRefunding { return domain.ErrOrderState }
if err := s.orders.Transition(txCtx, order.OrderNo,
    domain.OrderRefunding, domain.OrderRefunded, order.Version); err != nil { return err }
payment, err := s.payments.GetByOrderNo(txCtx, refund.OrderNo)
if err != nil { return err }
if err := s.payments.Transition(txCtx, payment.TransactionNo,
    domain.PaymentSuccess, domain.PaymentRefunded, payment.Version); err != nil { return err }
if err := s.locks.ReleaseByOrderNo(txCtx, refund.OrderNo, domain.SeatLockReleased); err != nil { return err }
if err := s.points.ReclaimOnRefund(txCtx, refund.UserID, refund.AmountCents, refund.RefundNo); err != nil { return err }
if err := s.box.Record(txCtx, domain.NewRefundEvent(order, refund.AmountCents, refund.RefundNo, time.Now())); err != nil { return err }
return s.refunds.MarkSuccess(txCtx, refund.RefundNo)''',
    '说明：退款回调的输入为 refund_no，成功回调重复到达时直接返回。订单、支付、座位锁、积分和票房事件在同一事务中完成状态迁移；任一写入失败都会回滚，保证一单一退和资金记录一致。'
)

# Language, punctuation, numbering, and factual consistency pass.
set_text(doc.paragraphs[21], '二〇二六年九月二日')
set_text(doc.paragraphs[25], '进一步理解和掌握 C、C++ 或 Java 等程序设计语言，熟悉控制语句、数组、面向对象编程、文件处理和数据库处理等概念，能够使用相应的集成开发环境（IDE）解决实际问题。')
set_text(doc.paragraphs[29], '当前主流程序设计语言包括 C、C++、Java 和 Python 等。请比较这些语言的主要特性，分析其在程序设计中的优势与不足，并结合选题选择合适的语言并说明理由。')
set_text(doc.paragraphs[31], '不同语言有各自的编码规范。请依照所选语言的规范编写代码，使程序具备规范、整洁、可读和易于维护等特点。')
set_text(doc.paragraphs[37], '每位同学应结合具体题目，独立完成程序设计实践，并在报告中说明实践过程中遇到的问题及个人体会。')
set_text(doc.paragraphs[39], '程序设计课程实践限定一周，请每位同学独立完成本次实践并按上述提纲撰写报告。报告不少于 8 页，正文使用小四号宋体，行距为单倍行距。')
set_text(doc.paragraphs[43], '编写程序，实现一个通用的电影院订票系统，主要功能模块如下：')
set_text(doc.paragraphs[58], '本项目后端采用简化的领域驱动设计（DDD），各层依赖关系清晰。')
set_text(doc.paragraphs[59], '一、基础业务：影片、影院、影厅、场次和优惠券等模块采用 Handler、Service 和 Repository 的分层结构。')
set_text(doc.paragraphs[60], 'Handler 负责接收并校验 HTTP 参数，Service 负责业务校验和事务控制，Repository 负责 PostgreSQL 数据读写。')
set_text(doc.paragraphs[61], '二、订单业务：用户提交座位后，后端先检查场次和座位状态，再在一个事务中写入座位锁、订单和订单明细。订单默认保留 15 分钟，超过时限仍未支付时，系统关闭订单并释放座位。')
set_text(doc.paragraphs[82], '（5）场次排片模块：场次将影片、影院、影厅、开场时间和票价关联起来。创建场次时需要检查影片与影厅状态，并避免同一影厅出现时间冲突；影院管理员可以调整票价或取消场次。用户端按照影片和影院查询可售场次，已取消或不可售的场次不作为正常购票入口。')
set_text(doc.paragraphs[75], '（11）会员积分与优惠券模块：支付成功后按照订单金额增加积分，并记录积分流水；退款时扣回本次订单产生的积分。用户可以查看当前积分和会员等级，也可以使用积分兑换优惠券。下单时，后端检查优惠券归属、有效期、使用门槛和状态，使用成功后将优惠券与订单绑定。')
set_text(doc.paragraphs[76], '（12）运营内容与数据看板模块：管理端可以维护首页 Banner、优惠券模板和发放记录。票房看板提供订单数、出票数、总票房、退款额和净票房，并支持按日期、影片和影院查看趋势及排行。看板数据来源于支付和退款产生的票房事件，超级管理员还可以执行对账，检查汇总结果是否一致。')
set_text(doc.paragraphs[77], '（10）票券核销模块：影院工作人员输入票券码后，系统查询对应订单、场次和座位信息；仅已支付且尚未核销的票券可以完成入场核销，重复核销或无效票券会被拒绝。')
doc.paragraphs[74]._p.addnext(doc.paragraphs[77]._p)
set_text(doc.paragraphs[87], '后端选择 Go 语言。C/C++ 需要开发者承担更多内存管理工作；Java 生态成熟，但运行环境和工程配置相对复杂；Rust 更适合强调内存安全和底层性能的系统开发；Python 开发效率较高，但在高并发服务中通常需要额外的性能优化。Go 语法简洁，具有良好的并发支持和部署便利性，其接口与组合机制也适合本项目的分层架构。')
set_text(doc.paragraphs[88], '前端采用 React 和 TypeScript。React 便于将页面拆分为电影卡片、场次列表、座位图、订单卡片和管理表格等可复用组件；TypeScript 通过静态类型约束接口参数和响应数据，有助于在开发阶段发现类型错误并提高代码可维护性。')
set_text(doc.paragraphs[104], '项目开发过程使用 Git 进行版本管理，各阶段提交记录均可追溯。')
set_text(doc.paragraphs[116], '5.1 核心请求链路')
set_text(doc.paragraphs[117], '用户提交选座请求后，Handler 完成身份与参数校验，Service 在事务中依次校验场次、抢占座位、创建订单及订单明细，并由 Repository 写入 PostgreSQL。支付回调沿相同分层进入支付服务，再联动订单、座位、优惠券、积分和票房状态。')
set_text(doc.paragraphs[125], '问题和讨论')
set_text(doc.paragraphs[126], '上线部署过程中曾出现影片图片加载不稳定的问题。初期使用外部图片地址，受网络延迟和资源可用性影响，部分图片无法及时显示；后续改为前端内置静态资源，降低了对外部对象存储和网络环境的依赖。')
set_text(doc.paragraphs[127], '首次在 macOS 环境中部署 Docker 时发现，Homebrew 安装的相关组件未自动包含 Docker Compose。完成组件安装并补充 docker-compose.yml 配置后，PostgreSQL、数据库迁移、后端和前端服务得以统一编排。')
set_text(doc.paragraphs[129], '本次课程实践提升了我的程序设计、需求分析和业务建模能力，使我能够独立完成从系统设计、编码实现、测试验证到部署运行的完整过程。通过订单、支付和座位锁等模块的实现，我进一步理解了状态机、事务、幂等和并发控制在实际业务系统中的作用。')
insert_after(doc.paragraphs[125], '项目初期的主要困难是需求分析、产品需求文档（PRD）和数据库表结构设计。通过查阅课程资料、技术文档和同类系统设计案例，逐步明确了业务边界、核心实体及其关联关系。完成总体架构和数据模型后，再按照模块拆分实现各项功能。')

# Remove obsolete image-only filler paragraphs where possible.
for idx in [113, 114, 115, 117]:
    if idx < len(doc.paragraphs) and not doc.paragraphs[idx].text.strip() and not doc.paragraphs[idx]._p.xpath('.//w:drawing'):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

# Keep the classification table together and repeat its header when WPS splits
# it at a page boundary.
for row in db_table.rows:
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit'); trPr.append(cant)
hdrPr = db_table.rows[0]._tr.get_or_add_trPr()
tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), 'true'); hdrPr.append(tblHeader)

# Academic wording polish in existing sections.
replacements = {
    'C /C++ 内存管理心智负担太重Web太难用 ；Java 太笨重了我不喜欢...（它更适合复杂规则场景：电商/金融）；Rust 不适合写业务，天生不支持方法嵌套连续调用。Python 我感觉性能太低了虽然底':
    'C/C++ 需要承担较多内存管理工作，Java 更适合规则复杂的大型业务，Rust 的生态与课程项目需求不完全匹配，Python 在高并发场景下需要额外的性能优化。因此后端选择 Go，以兼顾编译性能、并发能力、部署便利性和 Web 开发效率。',
    '因为我并没有购买 COS 或者 OSS 存储桶，用的是国外的展示 URL，服务器经常超时不显示，选择前端内置图片展示的方式，虽然丑是丑了点但是很不错。':
    '部署初期使用外部图片 URL，受网络延迟影响出现资源加载不稳定。后续改为前端内置静态资源，降低了对外部对象存储和跨境网络的依赖。',
}
for p in doc.paragraphs:
    if p.text in replacements:
        set_text(p, replacements[p.text])

# Add high-value screenshots from the running deployment. Code screenshots stay
# removed because the report now contains editable source code.
insert_figure_after(
    find_paragraph('（2）首页、推荐与搜索模块'),
    f'{ASSET_DIR}/01-user-homepage.png',
    '图2-4 用户端影院热映与场次聚合界面',
    '用户端首页将当前影院、热映影片、推荐内容和可购场次集中展示，运行结果如图2-4所示。'
)
insert_figure_after(
    find_paragraph('（3）影片管理模块'),
    f'{ASSET_DIR}/09-platform-movie-management.png',
    '图2-5 平台影片信息管理界面',
    '平台管理员可维护影片资料并控制影片上下架状态，管理界面如图2-5所示。'
)
insert_figure_after(
    find_paragraph('（4）影院、影厅与座位模块'),
    f'{ASSET_DIR}/07-cinema-hall-management.png',
    '图2-6 影院影厅与座位布局管理界面',
    '影厅管理界面将影院、影厅和座位布局关联起来，运行结果如图2-6所示。'
)
insert_figure_after(
    find_paragraph('（5）场次排片模块'),
    f'{ASSET_DIR}/06-cinema-session-scheduling.png',
    '图2-7 影院场次排片与票价管理界面',
    '影院运营人员可按授权影院维护场次、影厅和分级票价，实际管理界面如图2-7所示。'
)
insert_figure_after(
    find_paragraph('典型链路为'),
    f'{ASSET_DIR}/02-role-login.png',
    '图2-3 多角色统一登录界面',
    '系统在统一登录入口区分观众、影院运营和平台管理身份，如图2-3所示。'
)
insert_figure_after(
    find_paragraph('2.3 各功能模块设计'),
    f'{ASSET_DIR}/diagram-roles.png',
    '图2-2 系统角色功能与权限边界图',
    '在展开各功能模块之前，先对系统角色、职责范围和权限约束进行归纳，如图2-2所示。'
)
insert_figure_after(
    find_paragraph('（6）实时选座模块'),
    f'{ASSET_DIR}/03-user-realtime-seat-map.png',
    '图2-8 实时座位状态与用户选座界面',
    '实时选座页同时呈现可选、已选、已锁和已售状态，并展示锁座时限与预计金额，如图2-8所示。'
)
insert_figure_after(
    find_paragraph('（7）订单处理模块'),
    f'{ASSET_DIR}/04-user-paid-order-detail.png',
    '图2-9 已支付订单与电子票详情界面',
    '订单详情展示订单状态、座位、金额和电子票信息，可用于核对支付后的订单流转结果，如图2-9所示。'
)
insert_figure_after(
    find_paragraph('（12）运营内容与数据看板模块'),
    f'{ASSET_DIR}/05-platform-box-office-dashboard.png',
    '图2-11 平台票房统计与经营数据看板',
    '票房看板对支付订单、出票数、退款额、净票房和排行进行汇总，运行结果如图2-11所示。'
)
insert_figure_after(
    find_paragraph('（10）票券核销模块'),
    f'{ASSET_DIR}/08-cinema-ticket-verification.png',
    '图2-10 影院票券查询与核销界面',
    '影院工作人员可按票券码查询订单和座位信息，并完成入场核销，界面如图2-10所示。'
)
insert_figure_after(
    find_paragraph('2.2 详细设计'),
    f'{ASSET_DIR}/diagram-architecture.png',
    '图2-1 系统分层架构与请求链路图',
    '系统采用前后端分离与分层设计，整体分层、主要组件和请求调用方向如图2-1所示。'
)
insert_figure_after(
    find_paragraph('2.4 数据库表设计'),
    f'{ASSET_DIR}/diagram-database.png',
    '图2-12 数据库表归类与核心数据流转图',
    '数据库按业务域组织，核心表及购票售后数据流转关系如图2-12所示。'
)
insert_figure_after(
    find_paragraph('座位锁状态机：'),
    f'{ASSET_DIR}/diagram-seatstate.png',
    '图5-1 座位锁状态迁移图',
    '座位锁在创建订单、支付成功、订单取消、超时和退款事件下的状态迁移如图5-1所示。'
)

# Final content cleanup after all inserted paragraphs have settled. Match by
# text rather than source indices so old image placeholders cannot shift edits.
formal_go = '后端选择 Go 语言。C/C++ 需要开发者承担更多内存管理工作；Java 生态成熟，但运行环境和工程配置相对复杂；Rust 更适合强调内存安全和底层性能的系统开发；Python 开发效率较高，但在高并发服务中通常需要额外的性能优化。Go 语法简洁，具有良好的并发支持和部署便利性，其接口与组合机制也适合本项目的分层架构。'
formal_frontend = '前端采用 React 和 TypeScript。React 便于将页面拆分为电影卡片、场次列表、座位图、订单卡片和管理表格等可复用组件；TypeScript 通过静态类型约束接口参数和响应数据，有助于在开发阶段发现类型错误并提高代码可维护性。'
set_text(find_paragraph(formal_go), '编程语言的选择')
set_text(find_paragraph('我最后选择 Go 作为后端语言'), formal_go)
set_text(find_paragraph('前端采用 React 和 TypeScript。选择 React'), formal_frontend)
extra_frontend = find_paragraph(formal_frontend)
extra_frontend._element.getparent().remove(extra_frontend._element)
set_text(find_paragraph('项目开发过程使用 Git 进行版本管理'), 'Server 层统一使用 Axios 实例，请求拦截器添加 JWT 和 X-Request-ID，响应拦截器处理统一返回格式和登录过期。页面不直接调用 Axios，而是通过自定义 Hook 获取数据。')
set_text(find_paragraph('后端架构图：'), '后端分层架构及依赖方向如图2-1所示。')
for obsolete in ['支付状态机：', '然后就是上线部署的过程中', '本次课程实践提升了我的程序设计']:
    p = find_paragraph(obsolete)
    p._element.getparent().remove(p._element)
set_text(find_paragraph('本次课程中我锻炼了我编码能力'), '本次课程实践提升了我的程序设计、需求分析和业务建模能力，使我能够独立完成从系统设计、编码实现、测试验证到部署运行的完整过程。通过订单、支付和座位锁等模块的实现，我进一步理解了状态机、事务、幂等和并发控制在实际业务系统中的作用。')
set_text(find_paragraph('TypeScript 不使用 any'), 'TypeScript 不使用 any，接口数据统一定义类型。')
set_text(find_paragraph('分析：电影院订票系统不只是普通的增删改查'), '分析：电影院订票系统不仅包含常规的数据增删改查操作，还需要支持用户浏览影片、选择影院与场次、实时选座、创建订单和支付购票等完整流程。管理员负责维护影片、影院、影厅、排片和票房数据，因此系统划分 USER、SUPER_ADMIN、CINEMA_ADMIN 和 FINANCE 等角色，并在前后端实施权限校验。订单、支付和座位锁具有明确的状态变化，其状态变更必须遵循预定义的迁移规则。')
set_text(find_paragraph('登录成功后由后端签发 JWT'), '（1）用户与权限模块：用户端提供注册、登录、找回密码和修改密码功能。登录成功后由后端签发 JWT，前端通过 Axios 请求拦截器自动写入 Authorization 请求头。管理端分为超级管理员、影院管理员和财务人员；前端根据角色展示菜单，后端中间件再次检查接口权限，避免仅依赖前端菜单隐藏实施访问控制。')
set_text(find_paragraph('取消场次时，系统必须释放相关锁座并通知用户'), '影院管理员只管理被授权影院，负责影厅与座位维护、排片创建与调整、票价配置、场次关闭以及销售情况查看。取消场次时，系统释放相关锁座并处理关联订单；影院管理员不得修改平台角色权限和财务流水。')
set_text(find_paragraph('财务人员围绕资金与经营数据工作'), '财务人员具有经营数据只读权限，可以查看订单数、出票数、总票房、退款额、净票房以及按日期、影片和影院生成的统计结果，但不能修改订单、退款和排片状态。普通用户可以完成注册与登录，浏览影片、影院和场次，实时选座、锁座购票、查询订单与电子票，并申请退票或改签、使用优惠券和查看会员积分。典型链路为“管理员录入影片→审核上架→关联影院排片；用户浏览影片→选择影院和场次→锁座→创建订单→支付回调确认；财务查看票房统计结果”。')
set_text(find_paragraph('处理已经占用的优惠券'), '（7）订单处理模块：锁座成功后，系统在同一事务中创建订单和订单明细，明细保存场次、座位号和成交价格。订单金额由座位价格减去优惠券金额计算，用户可以在订单列表和详情页查看状态。待支付订单保留 15 分钟，超时任务关闭订单、释放座位，并解除已锁定优惠券与订单的绑定。')
for duplicate in ['本项目后端采用贫血 DDD 领域驱动设计', '一、基础业务：影片、影院、影厅、场次和优惠券等模块采用 Handler、Service']:
    p = find_paragraph(duplicate)
    p._element.getparent().remove(p._element)
detail_anchor = find_paragraph('二、订单业务：')
detail_anchor = insert_after(detail_anchor, '三、支付业务：创建支付流水时使用业务单号唯一约束，支付回调使用 event_id 唯一约束，防止同一笔订单被重复支付、重复出票或重复累计积分。')
insert_after(detail_anchor, '四、并发和定时任务：选座页每 5 秒刷新一次座位状态，但最终状态以后端事务处理结果为准。数据库在 seat_locks 表上建立部分唯一索引，保证同一场次的同一座位只能存在一条 LOCKED 或 BOOKED 状态的记录。后台定时任务负责关闭超时订单、释放座位以及重试失败的 Mock 支付回调。')
duplicate_session = find_paragraph('（5）场次排片模块：场次将影片')
duplicate_session._element.getparent().remove(duplicate_session._element)
set_text(find_paragraph('（5）场次排片模块：场次把影片'), '（5）场次排片模块：场次将影片、影院、影厅、开场时间和票价关联起来。创建场次时需要检查影片与影厅状态，并避免同一影厅出现时间冲突；平台管理员和影院管理员可以调整票价或取消场次。用户端按照影片和影院查询可售场次，已取消或不可售的场次不作为正常购票入口。')
set_text(find_paragraph('前端路由与接口调用模块：'), '（13）前端路由与接口调用模块：前端路由分为公共页面、用户页面和管理页面，管理页面再根据角色控制可访问菜单，并使用懒加载减少首次加载内容。Server 层只负责封装 Axios 接口，Feature Hook 使用统一的查询键、缓存和刷新规则，页面不直接拼接请求。请求拦截器负责 JWT 和 X-Request-ID，响应拦截器统一处理错误；搜索输入使用 300 毫秒防抖，提交按钮在请求过程中禁用，支付等关键请求另外携带幂等键。')
for old, new in [
    ('题目的要求', '一、题目的要求'),
    ('题目的设计过程', '二、题目的设计过程'),
    ('编程语言的选择', '三、编程语言的选择'),
    ('代码规范的说明', '四、代码规范的说明'),
    ('系统测试与运行结果', '六、系统测试与运行结果'),
    ('问题和讨论', '七、问题和讨论'),
    ('八、总结体会', '八、总结与体会'),
]:
    set_text(find_paragraph(old), new)
test_anchor = find_paragraph('前端执行 npm run build')
test_anchor = insert_after(test_anchor, '接口测试重点验证重复请求的处理结果。创建支付时携带 Idempotency-Key，payment_transactions 表使用业务类型和业务单号的联合唯一约束，支付回调使用 event_id 唯一约束。因此，重复调用不会生成第二笔交易，也不会重复出票或增加积分。')
insert_after(test_anchor, '部署方面已编写前端和后端 Dockerfile，并使用 docker-compose.yml 编排 PostgreSQL、数据库迁移、后端和前端服务。迁移容器在启动时执行表结构迁移和演示数据初始化，使系统能够在服务器环境中完成统一部署与运行验证。')

# Place the request-chain subsection before code and state-machine details.
core_heading = find_paragraph('5.1 核心请求链路')

# Keep headings and the source block with the content they introduce.
for heading in ['2.2 详细设计', '2.3 各功能模块设计', '2.4 数据库表设计',
                '5.1 核心请求链路',
                '座位锁状态机：', '5.1 核心请求链路', '系统测试与运行结果',
                '七、问题和讨论', '八、总结与体会']:
    find_paragraph(heading).paragraph_format.keep_with_next = True
find_paragraph('if order.PaidCents').paragraph_format.keep_together = True

# Avoid orphaned heading/figure groups and excessive trailing whitespace.
for p in doc.paragraphs:
    if p.text.strip().startswith(('图2-', '图5-')):
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.keep_together = True
    pf = p.paragraph_format
    if p.text.strip():
        pf.widow_control = True

# Ensure the payment callback excerpt remains present after paragraph insertions.
if not any('代码5-1 支付回调' in p.text for p in doc.paragraphs):
    add_code_excerpt(
        '五、核心代码的实现',
        '代码5-1 支付回调幂等与出票（payment_svc.go，源码节选）',
        '''callback, err := s.callbacks.InsertIfAbsent(txCtx, in.EventID)
if err != nil { return err }
if callback.ProcessedAt != nil { return nil }
payment, err := s.payments.GetByTransactionNo(txCtx, in.TransactionNo)
if err != nil { return err }
if payment.AmountCents != in.AmountCents { return domain.ErrPaymentAmountMismatch }
order, err := s.orders.GetOrderForUpdate(txCtx, payment.OrderNo)
if err != nil || order.Status != domain.OrderPendingPayment { return domain.ErrOrderState }
if err := s.payments.Transition(txCtx, payment.TransactionNo,
    domain.PaymentPending, domain.PaymentSuccess, payment.Version); err != nil { return err }
if err := s.orders.Transition(txCtx, order.OrderNo,
    domain.OrderPendingPayment, domain.OrderPaid, order.Version); err != nil { return err }
if err := s.locks.MarkBookedByOrderNo(txCtx, order.OrderNo); err != nil { return err }
if err := s.coupons.MarkUsedByOrderNo(txCtx, order.OrderNo); err != nil { return err }
if err := s.points.GrantOnPaid(txCtx, order.UserID, order.PaidCents, order.OrderNo); err != nil { return err }
if err := s.box.Record(txCtx, domain.NewPaidEvent(order, order.OrderNo, time.Now())); err != nil { return err }
return s.callbacks.MarkProcessed(txCtx, in.EventID)''',
        '说明：以 event_id 作为回调幂等键，重复通知直接返回；金额和旧状态校验通过后，支付、订单、座位、优惠券、积分和票房在同一事务内联动。'
    )

# Reconcile chapter order and remove superseded duplicate test/deployment drafts.
def move_before(src_text, dst_text):
    src = find_paragraph(src_text); dst = find_paragraph(dst_text)
    dst._p.addprevious(src._p)
move_before('六、系统测试与运行结果', '七、问题和讨论')
# Keep the test paragraphs with chapter six instead of leaving them after chapter seven.
for prefix in ('完成主要功能后，我先执行', '前端执行 npm run build', '接口测试重点验证重复请求', '部署方面已编写前端和后端 Dockerfile'):
    for p in list(doc.paragraphs):
        if p.text.startswith(prefix):
            find_paragraph('七、问题和讨论')._p.addprevious(p._p)
            break
for p in list(doc.paragraphs):
    if p.text.startswith(('接口测试中重点检查了重复请求', '部署方面已经编写后端和前端 Dockerfile', '问题和讨论\n', '还有就是第一次在 MacOS')):
        p._element.getparent().remove(p._element)
# Remove an obsolete duplicate left by the original template.
for p in list(doc.paragraphs):
    if p.text.startswith('四、并发和定时任务：选座页每 5 秒刷新一次座位状态，但是'):
        p._element.getparent().remove(p._element)
# Remove empty runs left between the final dashboard figure and the next module.
for p in list(doc.paragraphs):
    if not p.text.strip() and not p._p.xpath('.//w:drawing'):
        prev = p._p.getprevious(); nxt = p._p.getnext()
        if prev is not None and nxt is not None:
            prev_text = ''.join(prev.itertext())
            next_text = ''.join(nxt.itertext())
            if '图2-11' in prev_text or '（13）前端路由' in next_text:
                p._element.getparent().remove(p._element)
# Keep the feature list continuous after removing an obsolete original item.
for p in doc.paragraphs:
    if p.text.startswith('（10）票券核销模块'):
        set_text(p, p.text.replace('（10）', '（9）', 1))
    elif p.text.startswith('（11）会员积分与优惠券模块'):
        set_text(p, p.text.replace('（11）', '（10）', 1))
    elif p.text.startswith('（12）运营内容与数据看板模块'):
        set_text(p, p.text.replace('（12）', '（11）', 1))
    elif p.text.startswith('（13）前端路由与接口调用模块'):
        set_text(p, p.text.replace('（13）', '（12）', 1))

# Add substantive runtime interpretation below feature captions to use page
# space for input, processing, output, and exception behavior.
figure_notes = {
    '图2-3 多角色统一登录界面': '运行结果：登录成功后返回 JWT 和角色信息，前端据此加载对应菜单；密码错误、账号禁用或角色无权限时，接口返回统一错误并保持当前页面状态。',
    '图2-4 用户端影院热映与场次聚合界面': '操作步骤：先选择影院，再按片名或类型筛选影片。搜索无结果时显示空状态；输出为影片详情和该影院可购场次。',
    '图2-5 平台影片信息管理界面': '输入包括片名、类型、时长、上映日期和海报地址。保存前校验必填字段，上架前检查资料完整性；下架影片不再出现在用户端购票入口。',
    '图2-6 影院影厅与座位布局管理界面': '后端校验座位布局 JSON 后同步更新 seats 表。已被订单引用的座位不能直接删除，只能标记停用，以保护历史订单和票券信息。',
    '图2-7 影院场次排片与票价管理界面': '创建场次时检查影片和影厅状态，并校验同一影厅的时间冲突；校验通过后输出可售场次，冲突时保留原数据并返回提示。',
    '图2-8 实时座位状态与用户选座界面': '页面每 5 秒刷新一次座位状态，提交时以后端事务结果为准。并发用户抢同一座位时，仅一个请求获得 LOCKED，其余请求收到占用提示。',
    '图2-9 已支付订单与电子票详情界面': '订单页展示订单号、座位、成交金额、支付时间和 ticket_no。待支付订单显示 15 分钟倒计时，超时自动关闭；支付成功后票券才可核销。',
    '图2-10 影院票券查询与核销界面': '核销输入为 ticket_no，系统返回订单、场次和座位信息并记录核销时间。已核销票券、无效票券或非本影院场次均拒绝重复操作。',
    '图2-11 平台票房统计与经营数据看板': '看板以支付事件累计总票房，以退款事件计算退款额和净票房，并按日期、影片、影院聚合订单数和出票数，便于经营对比。',
}
for caption, note in figure_notes.items():
    cap = find_paragraph(caption)
    following = cap._p.getnext()
    if following is None or note not in ''.join(following.itertext()):
        insert_after(cap, note)

# Do not strand code headings at the bottom of a page. Let WPS flow the
# editable block naturally with its explanation instead of reserving a whole
# page for a keep-together group.
for p in doc.paragraphs:
    if p.text.startswith('代码5-'):
        p.paragraph_format.page_break_before = False
        p.paragraph_format.keep_with_next = True
    if p.text.startswith(('return s.tx.Run', 'if order.PaidCents', 'expired, err :=', 'if refund.Status')):
        p.paragraph_format.keep_together = False

# Collapse only redundant empty paragraphs in the functional section.
for p in list(doc.paragraphs):
    if not p.text.strip() and not p._p.xpath('.//w:drawing'):
        prev = p._p.getprevious(); nxt = p._p.getnext()
        if prev is not None and nxt is not None:
            prev_text = ''.join(prev.itertext())
            next_text = ''.join(nxt.itertext())
            if ('图2-' in prev_text or '图5-' in prev_text) and next_text:
                p._element.getparent().remove(p._element)

# Reattach each code caption directly before its own source paragraph. This
# prevents WPS from displaying a list of captions followed by unrelated code.
code_pairs = [
    ('代码5-2 锁座与创建订单', 'return s.tx.Run'),
    ('代码5-3 超时订单释放', 'expired, err :='),
    ('代码5-4 退款回调与资源回收', 'if refund.Status'),
]
for title_text, code_prefix in code_pairs:
    title = next((p for p in doc.paragraphs if p.text.startswith(title_text)), None)
    code = next((p for p in doc.paragraphs if p.text.startswith(code_prefix)), None)
    if title is not None and code is not None:
        code._p.addprevious(title._p)

doc.save(OUT)
print(OUT)
